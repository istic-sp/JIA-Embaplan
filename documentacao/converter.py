#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Converte documentacao/DOCUMENTACAO.md em um .docx com identidade visual SENAI.
Uso: python documentacao/converter.py
"""
import os
import re
import shutil
import subprocess
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE, "DOCUMENTACAO.md")
OUT_PATH = os.path.join(BASE, "Embaplan-Documentacao.docx")
ASSETS = os.path.join(BASE, "assets")

# ---------- Paleta SENAI ----------
ACCENT = RGBColor(0xE3, 0x06, 0x13)
ACCENT2 = RGBColor(0x8B, 0x04, 0x10)
INK = RGBColor(0x0B, 0x0B, 0x0B)
INK2 = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
RULE = "D1D5DB"
ZEBRA = "F9F7F7"
CELLHDR = "FBE9EB"
CODEBG = "0B0B0B"
CODEFG = RGBColor(0xE5, 0xE7, 0xEB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CALLOUTS = {
    "NOTE": ("1D4ED8", "EAF0FE", "ℹ", "Nota"),
    "TIP": ("047857", "E7F6F0", "✔", "Dica"),
    "WARNING": ("B45309", "FCF3E6", "⚠", "Atenção"),
    "DANGER": ("B91C1C", "FBE9E9", "⛔", "Risco"),
    "IMPORTANT": ("8B0410", "FBE9EB", "★", "Importante"),
}

FONT = "Calibri"
MONO = "Consolas"


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_borders(cell, color=RULE, sz=4, sides=("top", "bottom", "left", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def shade_paragraph(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def para_bottom_border(p, color, sz=12):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def left_bar(cell, color, sz=36):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    el = OxmlElement("w:left")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    borders.append(el)
    tcPr.append(borders)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(paragraph, text, base_color=INK2, base_size=10.5, base_bold=False):
    """Renderiza markdown inline (negrito/itálico/código)."""
    text = text.replace("\\|", "|")
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], base_color, base_size, base_bold)
        tok = m.group(0)
        if tok.startswith("**"):
            _run(paragraph, tok[2:-2], base_color, base_size, True)
        elif tok.startswith("`"):
            r = _run(paragraph, tok[1:-1], ACCENT2, base_size, False)
            r.font.name = MONO
        elif tok.startswith("*"):
            r = _run(paragraph, tok[1:-1], base_color, base_size, base_bold)
            r.italic = True
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], base_color, base_size, base_bold)


def _run(paragraph, text, color, size, bold):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    return r


# ---------- Mermaid ----------
def find_mmdc():
    for cand in ("mmdc", "mmdc.cmd"):
        if shutil.which(cand):
            return shutil.which(cand)
    return None


def render_mermaid(code, idx):
    mmdc = find_mmdc()
    if not mmdc:
        return None
    os.makedirs(ASSETS, exist_ok=True)
    mmd = os.path.join(ASSETS, f"diagram_{idx}.mmd")
    png = os.path.join(ASSETS, f"diagram_{idx}.png")
    with open(mmd, "w", encoding="utf-8") as f:
        f.write(code)
    try:
        subprocess.run(
            [mmdc, "-i", mmd, "-o", png, "-b", "white", "-s", "2"],
            check=True, capture_output=True, timeout=120,
        )
        return png if os.path.exists(png) else None
    except Exception:
        return None


# ---------- Construção do documento ----------
def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK2


def add_cover(doc, title, subtitle, meta):
    # Faixa superior vermelha
    band = doc.add_paragraph()
    band.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shade_paragraph(band, "E30613")
    r = band.add_run("SENAI · Documentação de Software")
    r.font.name = FONT
    r.font.size = Pt(12)
    r.font.color.rgb = WHITE
    r.bold = True

    for _ in range(4):
        doc.add_paragraph()

    pt = doc.add_paragraph()
    rt = pt.add_run(title)
    rt.font.name = FONT
    rt.font.size = Pt(34)
    rt.font.color.rgb = ACCENT
    rt.bold = True

    ps = doc.add_paragraph()
    rs = ps.add_run(subtitle)
    rs.font.name = FONT
    rs.font.size = Pt(14)
    rs.font.color.rgb = INK

    # Régua
    rule_p = doc.add_paragraph()
    para_bottom_border(rule_p, "E30613", sz=18)

    for _ in range(10):
        doc.add_paragraph()

    for line in meta:
        pm = doc.add_paragraph()
        rm = pm.add_run(line)
        rm.font.name = FONT
        rm.font.size = Pt(11)
        rm.font.color.rgb = GRAY

    foot = doc.add_paragraph()
    shade_paragraph(foot, "8B0410")
    rf = foot.add_run("Confidencial — uso interno")
    rf.font.name = FONT
    rf.font.size = Pt(9)
    rf.font.color.rgb = WHITE
    doc.add_page_break()


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(fldEnd)


def add_toc(doc):
    h = doc.add_paragraph()
    rh = h.add_run("Sumário")
    rh.font.name = FONT
    rh.font.size = Pt(18)
    rh.font.color.rgb = ACCENT
    rh.bold = True
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph()
    rn = note.add_run("(Clique com o botão direito → Atualizar campo para gerar o índice.)")
    rn.font.name = FONT
    rn.font.size = Pt(8.5)
    rn.italic = True
    rn.font.color.rgb = GRAY
    doc.add_page_break()


def setup_header_footer(doc, doc_title, project):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    header = sec.header
    hp = header.paragraphs[0]
    hp.text = ""
    tab = hp.paragraph_format
    # cria tabela de 1 linha 2 colunas no header para alinhar esquerda/direita
    htab = header.add_table(1, 2, sec.page_width - sec.left_margin - sec.right_margin)
    htab.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc = htab.cell(0, 0).paragraphs[0]
    rl = lc.add_run(doc_title)
    rl.font.name = FONT
    rl.font.size = Pt(8)
    rl.font.color.rgb = GRAY
    rc = htab.cell(0, 1).paragraphs[0]
    rc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rc.add_run(project)
    rr.font.name = FONT
    rr.font.size = Pt(8)
    rr.font.color.rgb = ACCENT2
    rr.bold = True

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Página ")
    r1.font.name = FONT
    r1.font.size = Pt(8)
    r1.font.color.rgb = GRAY
    add_field(fp, "PAGE")
    r2 = fp.add_run(" de ")
    r2.font.name = FONT
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY
    add_field(fp, "NUMPAGES")


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {min(level,4)}"]
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text.upper() if level == 1 else text)
    r.font.name = FONT
    r.bold = True
    if level == 1:
        r.font.size = Pt(18)
        r.font.color.rgb = ACCENT
    elif level == 2:
        r.font.size = Pt(14)
        r.font.color.rgb = INK
        para_bottom_border(p, RULE, sz=8)
    elif level == 3:
        r.font.size = Pt(12)
        r.font.color.rgb = INK2
    else:
        r.font.size = Pt(11)
        r.font.color.rgb = INK2


def add_paragraph_md(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text)


def add_bullet(doc, text, numbered=False, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    if numbered:
        badge = p.add_run(f" {number} ")
        badge.font.name = FONT
        badge.font.size = Pt(10)
        badge.bold = True
        badge.font.color.rgb = WHITE
        _shade_run(badge, "E30613")
        p.add_run("  ")
    else:
        b = p.add_run("• ")
        b.font.color.rgb = ACCENT
        b.bold = True
    add_runs(p, text)


def _shade_run(run, hexcolor):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    rPr.append(shd)


def add_code_block(doc, lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, CODEBG)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    cell.paragraphs[0].text = ""
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(ln if ln else " ")
        r.font.name = MONO
        r.font.size = Pt(9)
        r.font.color.rgb = CODEFG
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_mermaid(doc, code, idx, caption):
    png = render_mermaid(code, idx)
    if png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(png, width=Inches(6.0))
        except Exception:
            run.add_picture(png)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = cap.add_run(caption)
        rc.font.name = FONT
        rc.font.size = Pt(8.5)
        rc.italic = True
        rc.font.color.rgb = GRAY
    else:
        add_code_block(doc, code.splitlines())
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = cap.add_run(caption + " (diagrama Mermaid)")
        rc.font.name = FONT
        rc.font.size = Pt(8.5)
        rc.italic = True
        rc.font.color.rgb = GRAY


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    # split em | que não esteja escapado
    parts = re.split(r"(?<!\\)\|", line)
    return [c.strip() for c in parts]


def add_table_md(doc, rows):
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]
    ncols = len(header)
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    # cabeçalho
    for j, htext in enumerate(header):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, "E30613")
        set_cell_borders(cell, RULE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_runs(p, htext, base_color=WHITE, base_size=9.5, base_bold=True)
    # corpo
    for i, brow in enumerate(body):
        cells = tbl.add_row().cells
        for j in range(ncols):
            cell = cells[j]
            txt = brow[j] if j < len(brow) else ""
            if i % 2 == 1:
                set_cell_bg(cell, ZEBRA)
            set_cell_borders(cell, RULE)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_runs(p, txt, base_color=INK2, base_size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_callout(doc, kind, lines):
    color, fill, icon, label = CALLOUTS.get(kind, CALLOUTS["NOTE"])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    left_bar(cell, color, sz=36)
    set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    rl = p0.add_run(f"{icon}  {label}")
    rl.font.name = FONT
    rl.font.size = Pt(10)
    rl.bold = True
    rl.font.color.rgb = RGBColor.from_string(color)
    for ln in lines:
        pp = cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        add_runs(pp, ln, base_color=INK2, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_heading_styles(doc):
    for lvl in range(1, 5):
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = FONT
        st.font.color.rgb = ACCENT if lvl == 1 else INK


def parse_and_build(doc, md):
    lines = md.split("\n")
    i = 0
    n = len(lines)
    mermaid_idx = 0
    num_counter = 0
    # pular bloco do título inicial (capa já tratada externamente)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # separador horizontal
        if stripped == "---":
            i += 1
            continue

        # cabeçalho
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # bloco de código / mermaid
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # pula fechamento
            if lang == "mermaid":
                mermaid_idx += 1
                add_mermaid(doc, "\n".join(block), mermaid_idx,
                            f"Figura {mermaid_idx}")
            else:
                add_code_block(doc, block)
            continue

        # callout > [!TIPO]
        cm = re.match(r"^>\s*\[!(\w+)\]\s*$", stripped)
        if cm:
            kind = cm.group(1).upper()
            i += 1
            body = []
            while i < n and lines[i].strip().startswith(">"):
                body.append(lines[i].strip().lstrip(">").strip())
                i += 1
            body = [b for b in body if b]
            add_callout(doc, kind, body)
            continue

        # tabela
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|?[\s:\-|]+\|?$", lines[i + 1].strip()):
            tbl_rows = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_rows.append(lines[i])
                i += 1
            add_table_md(doc, tbl_rows)
            continue

        # lista numerada
        nm = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if nm:
            add_bullet(doc, nm.group(2), numbered=True, number=nm.group(1))
            i += 1
            continue

        # bullet
        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
            i += 1
            continue

        # parágrafo vazio
        if stripped == "":
            i += 1
            continue

        # parágrafo normal
        add_paragraph_md(doc, stripped)
        i += 1


def main():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md = f.read()

    # extrai título e subtítulo do topo, e remove o bloco de capa do corpo
    title = "Embaplan"
    subtitle = "Documentação Técnica e de Negócio"
    # corta tudo antes do primeiro "## Sumário Executivo" para o corpo
    body_start = md.find("## Sumário Executivo")
    body_md = md[body_start:] if body_start != -1 else md

    doc = Document()
    style_base(doc)
    configure_heading_styles(doc)

    sec = doc.sections[0]
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.8)

    add_cover(
        doc,
        "Embaplan",
        "Documentação Técnica e de Negócio",
        [
            "Plataforma de inteligência para anúncios de marketplace (Shopee Ads)",
            "Versão do documento: 1.0",
            "Data: 24/06/2026",
            "Autor: Doc Master",
        ],
    )
    setup_header_footer(doc, "Documentação Técnica e de Negócio", "Embaplan")
    add_toc(doc)
    parse_and_build(doc, body_md)

    doc.save(OUT_PATH)
    print(f"OK: {OUT_PATH}")


if __name__ == "__main__":
    main()
